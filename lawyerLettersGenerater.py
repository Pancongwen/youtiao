# -*- coding: utf-8 -*-
"""
Description: 
Mantainer: peter.pan
Mail: peter.pancongwen@gmail.com
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook


class NoticeLetterofLawyerParticipation():
    def __init__(
        self,
        date,               # 日期
        court_name,          # 法院名称
        lawsuit_number,      # 案件号
        lawsuit_content,     # 案件内容
        client_name,         # 委托人
        defendant_name,      # 被诉人姓名
        defendant_ID,        # 被诉人身份证号
        law_firm_name,        # 律所名
        lawyer_name          # 受委托律师姓名
    ):
        self.year = date[:4]
        self.month = date[4:6]
        self.day = date[6:]
        self.court_name = court_name
        self.lawsuit_number = lawsuit_number
        self.lawsuit_content = lawsuit_content
        self.client_name = client_name
        self.defendant_name = defendant_name
        self.defendant_ID = defendant_ID
        self.law_firm_name = law_firm_name
        self.lawyer_name = lawyer_name

    def _write_header(self, doc):
        doc_paragraph = doc.add_paragraph()
        header_run = doc_paragraph.add_run("律师参加诉讼通知函")
        header_run.bold = True
        header_run.font.size = Pt(24)
        header_run.font.name = '黑体'
        doc_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _write_lawsuit_number(self, doc, year, lawsuit_number):
        doc_paragraph = doc.add_paragraph()
        lawsuit_number_run = doc_paragraph.add_run("[" + year + "]年第 " + lawsuit_number + " 号")
        lawsuit_number_run.font.size = Pt(14)
        lawsuit_number_run.font.name = '黑体'
        doc_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _write_content(self, doc, lawsuit_content, client_name, defendant_name, defendant_ID, lawyer_name):
        content_style = doc.styles['Normal']
        content_style.font.name = '宋体'
        content_style.font.size = Pt(14)
        doc_paragraph = doc.add_paragraph()
        doc_paragraph.style = doc.styles['Normal']
        doc_paragraph_format = doc_paragraph.paragraph_format
        doc_paragraph_format.first_line_indent = Inches(0.5)
        doc_paragraph_format.line_spacing = 2
        doc_paragraph_format.space_before = Pt(40)

        doc_paragraph.add_run("贵院受理的 ")
        doc_paragraph.add_run(client_name).underline = True
        doc_paragraph.add_run(" 与 ")
        doc_paragraph.add_run(defendant_name + "(身份证号：" + defendant_ID + ")").underline = True
        doc_paragraph.add_run(lawsuit_content).underline = True
        doc_paragraph.add_run("一案（诉讼），现由 ")
        doc_paragraph.add_run(client_name).underline = True
        doc_paragraph.add_run(" 委托本所 ")
        doc_paragraph.add_run(lawyer_name).underline = True
        doc_paragraph.add_run(" 律师为其诉讼代理人，特此通告。")
        doc_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _write_close(self, doc, court_name):
        content_style = doc.styles['Normal']
        content_style.font.name = '宋体'
        content_style.font.size = Pt(14)
        doc_paragraph = doc.add_paragraph()
        doc_paragraph.style = doc.styles['Normal']
        doc_paragraph_format = doc_paragraph.paragraph_format
        doc_paragraph_format.first_line_indent = Inches(0.5)
        doc_paragraph_format.space_before = Pt(40)
        doc_paragraph.add_run("此致")
        doc_paragraph = doc.add_paragraph()
        doc_paragraph.style = doc.styles['Normal']
        doc_paragraph_format = doc_paragraph.paragraph_format
        doc_paragraph_format.first_line_indent = Inches(0.5)
        doc_paragraph_format.space_before = Pt(40)
        doc_paragraph.add_run(court_name).underline = True

    def _write_signature(self, doc, law_firm_name, year, month, day):
        content_style = doc.styles['Normal']
        content_style.font.name = '宋体'
        content_style.font.size = Pt(14)
        doc_paragraph = doc.add_paragraph()
        doc_paragraph.style = doc.styles['Normal']
        doc_paragraph.add_run(law_firm_name)
        doc_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc_paragraph = doc.add_paragraph()
        doc_paragraph.style = doc.styles['Normal']
        doc_paragraph.add_run(year + " 年" + month + " 月" + day + " 日")
        doc_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def write_doc(self, doc, filepath):
        self.doc = doc
        self._write_header(self.doc)
        self._write_lawsuit_number(self.doc, self.year, self.lawsuit_number)
        self._write_content(self.doc, self.lawsuit_content, self.client_name, self.defendant_name, self.defendant_ID, self.lawyer_name)
        self._write_close(self.doc, self.court_name)
        self._write_signature(self.doc, self.law_firm_name, self.year, self.month, self.day)
        doc.save(filepath+self.defendant_name+"_"+self.defendant_ID+".docx")

def read_xlsx(infofile):
    wb = load_workbook(filename = infofile)
    sheet = wb['info']
    return sheet

def get_last_defendant_line_number(sheet):
    line_number = 10
    while True:
        defendant_name_cell = 'A' + str(line_number)
        if sheet[defendant_name_cell].value is not None:
            line_number += 1
        else:
            break
    return line_number

if __name__ == "__main__":    
    info = read_xlsx('info.xlsx')
    date = str(info['B1'].value)
    lawsuit_number = str(info['B2'].value)
    lawsuit_content = info['B3'].value
    client_name = info['B4'].value
    law_firm_name = info['B5'].value
    lawyer_name = info['B6'].value
    court_name = info['B7'].value

    last_defendant_line_number = get_last_defendant_line_number(info)

    for defendant_line in range(10, last_defendant_line_number):
        defendant_name = info['A'+str(defendant_line)].value
        defendant_ID = info['B'+str(defendant_line)].value
        lawsuit_number += str(defendant_line -9)

        doc = Document()
        letter_doc = NoticeLetterofLawyerParticipation(
            date,
            court_name,
            lawsuit_number,
            lawsuit_content,
            client_name,
            defendant_name,
            defendant_ID,
            law_firm_name,
            lawyer_name
        )
        letter_doc.write_doc(doc, "./docs/")
